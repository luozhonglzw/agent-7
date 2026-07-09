"""DOCX document parser using python-docx.

Extracts paragraphs, headings, and tables from Word documents.

Usage::

    from app.core.parsers.docx_parser import DOCXParser

    parser = DOCXParser()
    doc = parser.parse(Path("report.docx"))
"""

import logging
from pathlib import Path

from app.core.parsers.base import BaseParser, ParsedDocument, ParsedPage, ParserError

logger = logging.getLogger(__name__)

# Map Word heading styles to heading levels.
_HEADING_STYLES = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
}


class DOCXParser(BaseParser):
    """DOCX parser backed by python-docx.

    Walks paragraphs in order, detects headings via Word styles,
    and extracts table content as TSV text.
    """

    def supported_extensions(self) -> list[str]:
        """Return supported extensions.

        Returns:
            List containing ``".docx"`` and ``".doc"``.
        """
        return [".docx", ".doc"]

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a DOCX file.

        Args:
            file_path: Absolute path to the DOCX file.

        Returns:
            Parsed document.

        Raises:
            ParserError: If python-docx cannot open the file.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ParserError(
                detail="python-docx is not installed",
                file_path=str(file_path),
            ) from exc

        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            raise ParserError(
                detail=f"Cannot open DOCX: {exc}", file_path=str(file_path)
            ) from exc

        pages: list[ParsedPage] = []
        full_text_parts: list[str] = []
        title = ""

        # Core properties → metadata
        props = doc.core_properties
        metadata = {}
        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["created"] = str(props.created)
        if props.title:
            title = props.title

        current_page = ParsedPage(page_number=1)
        page_counter = 1

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                paragraph = None
                for p in doc.paragraphs:
                    if p._element is element:
                        paragraph = p
                        break
                if paragraph is None:
                    continue

                text = paragraph.text.strip()
                if not text:
                    continue

                style_name = paragraph.style.name if paragraph.style else ""

                # Detect heading
                heading_level = _HEADING_STYLES.get(style_name)
                if heading_level:
                    # Flush current page if it has content
                    if current_page.content.strip():
                        pages.append(current_page)
                        full_text_parts.append(current_page.content)
                        page_counter += 1
                        current_page = ParsedPage(page_number=page_counter)

                    current_page.heading = text
                    current_page.heading_level = heading_level

                    # First heading becomes title if not set
                    if not title:
                        title = text

                    continue

                # Normal paragraph
                current_page.content += text + "\n"

            elif tag == "tbl":
                # Table — extract as TSV
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                if table is None:
                    continue

                table_text = self._extract_table(table)
                if table_text:
                    current_page.content += table_text + "\n"

        # Flush last page
        if current_page.content.strip():
            pages.append(current_page)
            full_text_parts.append(current_page.content)

        if not title:
            title = file_path.stem

        full_text = "\n\n".join(full_text_parts)

        logger.debug(
            "Parsed DOCX %s: %d pages, %d chars",
            file_path.name, len(pages), len(full_text),
        )

        return ParsedDocument(
            title=title,
            pages=pages,
            full_text=full_text,
            metadata=metadata,
        )

    @staticmethod
    def _extract_table(table: object) -> str:
        """Convert a DOCX table to TSV text.

        Args:
            table: python-docx Table object.

        Returns:
            Tab-separated table content.
        """
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\t", " ") for cell in row.cells]
            rows.append("\t".join(cells))
        return "\n".join(rows)
